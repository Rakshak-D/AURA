from typing import BinaryIO
import logging

logger = logging.getLogger(__name__)

def parse_document(file: BinaryIO, content_type: str) -> str:
    """Parse document and extract text with proper error handling"""
    try:
        if content_type == 'text/plain' or content_type == 'text/markdown':
            try:
                return file.read().decode('utf-8')
            except UnicodeDecodeError:
                return file.read().decode('latin-1')
        
        elif content_type == 'application/pdf':
            try:
                import PyPDF2
                file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page.extract_text()
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
                        continue
                
                return text if text.strip() else "Could not extract text from PDF"
            
            except ImportError:
                logger.error("PyPDF2 not installed")
                return "ERROR: PDF parsing requires PyPDF2. Install with: pip install PyPDF2"
            except Exception as e:
                logger.exception(f"PDF parsing error: {e}")
                return f"Error parsing PDF: {str(e)}"
        
        elif content_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            try:
                from docx import Document
                file.seek(0)
                doc = Document(file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + " | ".join([cell.text for cell in row.cells])
                
                return text if text.strip() else "No content found in document"
            
            except ImportError:
                logger.error("python-docx not installed")
                return "ERROR: DOCX parsing requires python-docx. Install with: pip install python-docx"
            except Exception as e:
                logger.exception(f"DOCX parsing error: {e}")
                return f"Error parsing DOCX: {str(e)}"
        
        else:
            logger.warning(f"Unsupported content type: {content_type}")
            try:
                return file.read().decode('utf-8')
            except UnicodeDecodeError:
                return file.read().decode('latin-1')
    
    except Exception as e:
        logger.exception(f"Document parsing error: {e}")
        return f"Error parsing document: {str(e)}"

def sanitize_filename(filename: str, max_length: int = 255) -> str:
    """Sanitize filename to prevent directory traversal attacks"""
    import os
    import re
    
    # Remove path separators
    filename = os.path.basename(filename)
    
    # Remove dangerous characters
    filename = re.sub(r'[^\w\s.-]', '', filename)
    
    # Limit length
    filename = filename[:max_length]
    
    # Ensure it's not empty
    if not filename:
        filename = "document"
    
    return filename